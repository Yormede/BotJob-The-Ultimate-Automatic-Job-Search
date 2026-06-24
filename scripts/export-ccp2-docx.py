from pathlib import Path
import re
import json
import os
import subprocess

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ccp2" / "dossier-projet-botjob.md"
OUTPUT = ROOT / "docs" / "ccp2" / "dossier-projet-botjob.docx"
MERMAID_DIR = ROOT / "docs" / "ccp2" / "assets" / "generated-mermaid"
NODE = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "node" / "bin" / "node.exe"
NODE_MODULES = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "node" / "node_modules"
PNPM_NODE_MODULES = NODE_MODULES / ".pnpm" / "node_modules"


def strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text.strip()


def add_table_as_text(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    headers = [strip_inline(cell) for cell in rows[0]]
    for row in rows[1:]:
        cells = [strip_inline(cell) for cell in row]
        if len(cells) == 2:
            text = f"{cells[0]} : {cells[1]}"
        elif len(cells) == 3:
            text = f"{cells[0]} - {cells[1]} : {cells[2]}"
        else:
            pairs = []
            for index, value in enumerate(cells):
                label = headers[index] if index < len(headers) else f"Champ {index + 1}"
                pairs.append(f"{label} : {value}")
            text = " ; ".join(pairs)
        doc.add_paragraph(text, style="List Bullet")
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)
        if not is_separator:
            rows.append(cells)
        i += 1
    return rows, i


def collect_mermaid_blocks(lines: list[str]) -> list[str]:
    blocks = []
    in_code = False
    lang = ""
    buf = []
    for line in lines:
        if line.startswith("```"):
            if not in_code:
                in_code = True
                lang = line[3:].strip()
                buf = []
            else:
                if lang == "mermaid":
                    blocks.append("\n".join(buf))
                in_code = False
            continue
        if in_code:
            buf.append(line)
    return blocks


def render_mermaid(lines: list[str]) -> dict[int, Path]:
    blocks = collect_mermaid_blocks(lines)
    if not blocks:
        return {}
    MERMAID_DIR.mkdir(parents=True, exist_ok=True)
    payload = [{"id": f"diagram-{i + 1:02d}", "code": code} for i, code in enumerate(blocks)]
    json_path = MERMAID_DIR / "diagrams.json"
    json_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    env = os.environ.copy()
    env["NODE_PATH"] = os.pathsep.join([str(NODE_MODULES), str(PNPM_NODE_MODULES)])
    try:
        subprocess.run(
            [str(NODE), str(ROOT / "scripts" / "render-mermaid.mjs"), str(json_path), str(MERMAID_DIR)],
            check=True,
            cwd=str(ROOT),
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=60,
        )
    except Exception as exc:
        print(f"Mermaid render skipped: {exc}")
        return {}
    return {i: MERMAID_DIR / f"diagram-{i + 1:02d}.png" for i in range(len(blocks))}


def export() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    mermaid_images = render_mermaid(lines)
    mermaid_index = 0
    in_frontmatter = False
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if i == 0 and line == "---":
            in_frontmatter = True
            i += 1
            continue
        if in_frontmatter:
            if line == "---":
                in_frontmatter = False
            i += 1
            continue

        if line.strip() == "<!-- pagebreak -->":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                if code_lang == "mermaid":
                    image_path = mermaid_images.get(mermaid_index)
                    mermaid_index += 1
                    if image_path and image_path.exists():
                        doc.add_picture(str(image_path), width=Inches(6.4))
                        doc.add_paragraph()
                    else:
                        para = doc.add_paragraph()
                        run = para.add_run("\n".join(code_lines))
                        run.font.name = "Consolas"
                        run.font.size = Pt(8.5)
                else:
                    para = doc.add_paragraph()
                    run = para.add_run("\n".join(code_lines))
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table_as_text(doc, rows)
            continue

        if line.startswith("# "):
            doc.add_heading(strip_inline(line[2:]), level=1)
        elif line.startswith("## "):
            doc.add_heading(strip_inline(line[3:]), level=2)
        elif line.startswith("### "):
            doc.add_heading(strip_inline(line[4:]), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(strip_inline(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(strip_inline(re.sub(r"^\d+\. ", "", line)), style="List Number")
        elif line.startswith("> "):
            doc.add_paragraph(strip_inline(line[2:]))
        elif (
            line.startswith("[CAPTURE A INSERER")
            or line.startswith("[CAPTURE A AJOUTER")
            or line.startswith("[PREUVE ")
        ):
            para = doc.add_paragraph()
            run = para.add_run(strip_inline(line))
            run.italic = True
        else:
            doc.add_paragraph(strip_inline(line))

        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    export()
